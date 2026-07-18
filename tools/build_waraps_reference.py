from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_URL = "https://api-docs.waraps.org/"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 102, 114)


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill="F4F6F9"):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("WARA Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("WARA-PS API Specifications - Consolidated Reference")
    set_run_font(header_run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Compiled from api-docs.waraps.org")
    set_run_font(footer_run, size=8.5, color=MUTED)


def add_inline(paragraph, text: str):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"\s*<br\s*/?>\s*", " ", text, flags=re.IGNORECASE)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_markdown(doc: Document, text: str):
    in_code = False
    code_lines: list[str] = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        for line in code_lines:
            p = doc.add_paragraph(style="WARA Code")
            p.paragraph_format.keep_together = True
            shade_paragraph(p)
            run = p.add_run(line if line else " ")
            set_run_font(run, name="Consolas", size=8)
        code_lines = []

    for raw_line in text.splitlines():
        line = re.sub(r"\s*<br\s*/?>\s*", " ", raw_line.rstrip(), flags=re.IGNORECASE)
        if line.startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip() or line.lstrip().startswith("<!--"):
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            title = re.sub(r"\*\*([^*]+)\*\*", r"\1", heading.group(2)).strip()
            doc.add_heading(title, level=level)
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            continue
        ordered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if ordered:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, ordered.group(1))
            continue
        if line.startswith("|") and line.endswith("|"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            add_inline(p, line.strip("| ").replace("|", "  |  "))
            continue
        p = doc.add_paragraph()
        add_inline(p, line)
    flush_code()


def main(source_dir: str, output_file: str):
    source = Path(source_dir)
    paths = [line.strip() for line in (source / "manifest.txt").read_text(encoding="utf-8").splitlines() if line.strip()]
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("WARA-PS API Specifications")
    set_run_font(run, size=30, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("Consolidated technical reference")
    set_run_font(run, size=15, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Source: {BASE_URL} | Retrieved: {date.today().isoformat()} | Pages collected: {len(paths)}")
    set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()

    doc.add_heading("About this reference", level=1)
    p = doc.add_paragraph()
    add_inline(p, "This document consolidates the publicly available Markdown pages linked from the WARA-PS API documentation site. It preserves the published explanations, topic conventions, JSON examples, ROS examples, task commands, Task Specification Tree material, and service-communication material in one navigable Word document.")
    p = doc.add_paragraph()
    add_inline(p, "The source remains authoritative. This compilation does not add protocol behavior beyond the published pages; each source page is identified with its direct URL.")

    doc.add_heading("Contents", level=1)
    groups = [
        ("Agent communication overview and topic conventions", "Agent levels, naming, heartbeat, sensor, execution information, and ROS-MQTT bridge."),
        ("L2 tasks and commands", "Command envelope, feedback, and individual task definitions."),
        ("L3-L4 and Task Specification Tree", "TST topics, commands, execution information, and feedback messages."),
        ("Service communication", "Resource pool, object detector, and U-Space Service Provider workflows."),
    ]
    for label, detail in groups:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, f"**{label}.** {detail}")

    current_group = None
    for relative_path in paths:
        parts = Path(relative_path).parts
        group = parts[0]
        if group != current_group:
            current_group = group
            label = "Agent Communication" if group == "agent_communication" else "Service Communication"
            doc.add_page_break()
            doc.add_heading(label, level=1)

        content = (source / relative_path).read_text(encoding="utf-8")
        page_title_match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        page_title = page_title_match.group(1).strip() if page_title_match else Path(relative_path).stem.replace("_", " ")
        doc.add_heading(page_title, level=2)
        source_line = doc.add_paragraph()
        source_line.paragraph_format.space_after = Pt(7)
        source_run = source_line.add_run(f"Source page: {BASE_URL}{relative_path}")
        set_run_font(source_run, size=8.5, color=MUTED, italic=True)
        content_without_title = re.sub(r"^#\s+.+$", "", content, count=1, flags=re.MULTILINE)
        add_markdown(doc, content_without_title)

    Path(output_file).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    print(f"Wrote {output_file}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
