#!/usr/bin/env python3
"""Build the SysCraft Arena domain and SIDC catalog Word document."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SysCraft_Arena_Domain_SIDC_Catalog.docx"
ICON_DIR = ROOT / "docs" / "assets" / "domain-catalog-icons"

ROWS = [
    ("Air", "air", "130301000011010000000000000000", "Fixed-wing aircraft", "Air symbol"),
    ("Ground", "ground", "130315000016040000000000000000", "Utility ground vehicle", "Land-equipment vehicle symbol"),
    ("Surface", "surface", "130330000012050100000000000000", "Patrol craft", "Sea-surface vessel symbol"),
    ("Space", "space", "130305000011070000000000000000", "Satellite / orbital platform", "Space satellite symbol"),
    ("Cyber", "cyber", "130360000014010000000000000000", "Cyber operations node", "Cyberspace core-router symbol"),
    ("Social", "social", "130327000011010100000000000000", "Operator / dismounted individual", "Dismounted-individual symbol"),
]


def set_font(run, *, size=11, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_cell_text(cell, text, *, size=10, bold=False, color="102A43", mono=False):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color, name="Courier New" if mono else "Calibri")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("SysCraft Arena Domain & SIDC Catalog")
    set_font(run, size=24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Reference examples for MQTT SensorInfo entity snapshots")
    set_font(run, size=12, color="52616B")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    run = note.add_run("How the Arena renders entities. ")
    set_font(run, size=10.5, bold=True, color="1F4D78")
    run = note.add_run("The domain controls the Assets grouping. The SIDC controls the MIL-STD symbol drawn on the map. Use a SIDC that describes the specific entity; the values below are the current Arena reference defaults.")
    set_font(run, size=10.5, color="334E68")

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Domain reference table")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [900, 1050, 2500, 2200, 2710]
    set_table_geometry(table, widths)
    headers = ["Icon", "Domain", "SIDC", "Example entity", "Catalog note"]
    for cell, label in zip(table.rows[0].cells, headers):
        shade(cell, "E8EEF5")
        add_cell_text(cell, label, size=10, bold=True, color="1F4D78")

    for domain, icon_key, sidc, entity, note in ROWS:
        row = table.add_row()
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        icon_cell = row.cells[0]
        icon_p = icon_cell.paragraphs[0]
        icon_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_p.paragraph_format.space_after = Pt(0)
        icon_p.add_run().add_picture(str(ICON_DIR / f"{icon_key}.png"), width=Inches(0.46))
        add_cell_text(row.cells[1], domain, size=10, bold=True)
        add_cell_text(row.cells[2], sidc, size=8.2, mono=True)
        add_cell_text(row.cells[3], entity, size=10)
        add_cell_text(row.cells[4], note, size=9.3, color="52616B")

    doc.add_paragraph()
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run("MQTT placement")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Put the selected SIDC in the full, low-rate SensorInfo packet: ")
    set_font(run, size=10.5, color="334E68")
    run = p.add_run('"arena-entity": { "domain": "ground", "sidc": "130315000016040000000000000000" }')
    set_font(run, size=9.2, color="1F4D78", name="Courier New")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Keep HeartBeat packets small; they identify and refresh the entity but do not need to repeat the SIDC.")
    set_font(run, size=10.5, color="334E68")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("SysCraft Arena MQTT reference")
    set_font(footer_run, size=8.5, color="6B7C93")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
